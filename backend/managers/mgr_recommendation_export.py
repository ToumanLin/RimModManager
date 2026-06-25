from __future__ import annotations

import html
import platform
import re
import shutil
import tempfile
from pathlib import Path
from typing import Any, TypeAlias
from urllib.parse import quote

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# Pillow 的默认字体和 TrueType 字体类型不同，统一成别名可以让图片排版函数保持清晰。
PillowFont: TypeAlias = ImageFont.ImageFont | ImageFont.FreeTypeFont


class RecommendationExportManager:
    """生成面向分享推荐的模组介绍文件。"""

    # 只有单文件导出格式需要扩展名；剪贴板和纯图片不走文件保存对话框。
    FORMAT_EXTENSIONS = {
        "txt": ".txt",
        "markdown": ".md",
        "docx": ".docx",
        "pdf": ".pdf",
    }
    DOCUMENT_FORMATS = {"txt", "markdown", "docx", "pdf"}
    ALL_FORMATS = {*DOCUMENT_FORMATS, "image", "clipboard"}
    # Windows 文件名限制最严格，统一按这个规则清理可以避免跨平台导出路径出错。
    INVALID_FILENAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]+')
    RESERVED_NAMES = {
        "con", "prn", "aux", "nul",
        *(f"com{i}" for i in range(1, 10)),
        *(f"lpt{i}" for i in range(1, 10)),
    }
    # Markdown 转义分两步：先去掉已有反斜杠，再按最终格式统一转义，避免出现双重转义。
    MARKDOWN_ESCAPE_RE = re.compile(r"([\\`*_{}\[\]()#+\-.!|>])")
    EXISTING_MARKDOWN_ESCAPE_RE = re.compile(r"\\+([\\`*_{}\[\]()#+\-.!|>])")

    def export(self, payload: dict[str, Any], target_path: str | None = None, target_dir: str | None = None) -> dict[str, Any]:
        # 所有格式先归一化数据，后续写文件只消费稳定字段，避免每个格式重复处理别名、标签和空值。
        normalized = self.normalize_payload(payload)
        export_format = normalized["format"]
        items = normalized["mods"]
        options = normalized["options"]

        if export_format == "clipboard":
            # 剪贴板只返回纯文本，不在后端写临时文件。
            return {
                "format": export_format,
                "text": self.render_plain_text(items, options),
                "count": len(items),
            }
        if export_format == "txt":
            return self._write_text_file(target_path, normalized["source_name"], items, options)
        if export_format == "markdown":
            return self._write_markdown_dir(target_dir, normalized["source_name"], items, options)
        if export_format == "docx":
            return self._write_docx(target_path, normalized["source_name"], items, options)
        if export_format == "pdf":
            return self._write_pdf(target_path, normalized["source_name"], items, options)
        if export_format == "image":
            return self._write_images(target_dir, normalized["source_name"], items, options)
        raise ValueError(f"不支持的推荐导出格式: {export_format}")

    def normalize_payload(self, payload: dict[str, Any] | None) -> dict[str, Any]:
        payload = payload or {}
        export_format = str(payload.get("format") or "txt").strip().lower()
        if export_format not in self.ALL_FORMATS:
            raise ValueError(f"不支持的推荐导出格式: {export_format}")
        # 导出名称会同时作为文档标题和默认文件名来源，所以先清理成可展示的单行文本。
        source_name = self._clean_inline_text(
            payload.get("source_name") or payload.get("group_name") or payload.get("group", {}).get("name") or "模组推荐"
        )
        options = self._normalize_options(payload.get("options") or {})
        # 跳过非 dict 项，保证前端选中列表里混入异常值时不会拖垮整次导出。
        mods = [
            self._normalize_mod(index, mod, options)
            for index, mod in enumerate(payload.get("mods") or [], start=1)
            if isinstance(mod, dict)
        ]
        if not mods:
            raise ValueError("没有可导出的模组")
        return {
            "format": export_format,
            "source_name": source_name,
            "options": options,
            "mods": mods,
        }

    def default_filename(self, payload: dict[str, Any] | None) -> str:
        normalized = self.normalize_payload(payload)
        export_format = normalized["format"]
        if export_format not in self.FORMAT_EXTENSIONS:
            return ""
        # 文件名只使用清理后的导出名称，避免用户自定义标题里的冒号、斜杠造成路径错误。
        base_name = self.sanitize_filename(normalized["source_name"], "模组推荐清单")
        return f"{base_name}{self.FORMAT_EXTENSIONS[export_format]}"

    def ensure_extension(self, path_value: str, export_format: str) -> str:
        extension = self.FORMAT_EXTENSIONS.get(str(export_format or "").strip().lower())
        if not extension:
            return str(path_value or "")
        path = Path(str(path_value or ""))
        if path.suffix.lower() == extension:
            return str(path)
        return str(path.with_suffix(extension))

    def file_types_for_format(self, export_format: str) -> tuple[str, ...]:
        normalized = str(export_format or "").strip().lower()
        return {
            "txt": ("Text Files (*.txt)", "All Files (*.*)"),
            "markdown": ("Markdown Files (*.md)", "All Files (*.*)"),
            "docx": ("Word Documents (*.docx)", "All Files (*.*)"),
            "pdf": ("PDF Files (*.pdf)", "All Files (*.*)"),
        }.get(normalized, ("All Files (*.*)",))

    @classmethod
    def sanitize_filename(cls, value: Any, fallback: str = "untitled", max_length: int = 96) -> str:
        text = str(value or "").strip()
        text = cls.INVALID_FILENAME_RE.sub("_", text)
        text = re.sub(r"\s+", " ", text).strip(" ._")
        if not text:
            text = fallback
        if text.lower() in cls.RESERVED_NAMES:
            text = f"{text}_"
        # 限制长度能避免 Windows 长路径和压缩/分享场景里的文件名过长问题。
        return text[:max_length].rstrip(" ._") or fallback

    def render_plain_text(self, items: list[dict[str, Any]], options: dict[str, Any], title: str | None = None) -> str:
        blocks = [self._render_text_block(item, options) for item in items]
        parts = [self._clean_inline_text(title)] if title else []
        parts.extend(block for block in blocks if block)
        return "\n\n".join(part for part in parts if part).strip() + "\n"

    def render_markdown(self, items: list[dict[str, Any]], options: dict[str, Any], image_names: dict[int, str] | None = None, title: str | None = None) -> str:
        blocks = []
        image_names = image_names or {}
        for item in items:
            item_title = item["display_name"]
            if options["include_sequence"]:
                item_title = f"{item['sequence']}. {item_title}"
            lines = [f"## {self._escape_markdown_text(item_title)}"]
            image_name = image_names.get(item["index"])
            if options["include_cover"] and image_name:
                # Markdown 图片路径固定指向同级 img 目录，文件名通过 quote 处理空格等 URL 字符。
                lines.append(f"![{self._markdown_alt_text(item['display_name'])}](./img/{quote(image_name)})")
            lines.extend(self._render_markdown_field_lines(item, options, include_name=False))
            blocks.append("\n\n".join(lines))
        body = "\n\n---\n\n".join(blocks)
        heading = self._clean_inline_text(title)
        if heading:
            return f"# {self._escape_markdown_text(heading)}\n\n{body}".strip() + "\n"
        return body.strip() + "\n"

    def _normalize_options(self, raw_options: dict[str, Any]) -> dict[str, Any]:
        # 选项默认偏向“推荐分享可读”，包名这类偏技术的信息默认关闭。
        return {
            "include_sequence": raw_options.get("include_sequence", True) is not False,
            "include_cover": raw_options.get("include_cover", True) is not False,
            "include_tags": raw_options.get("include_tags", True) is not False,
            "include_group_names": raw_options.get("include_group_names", True) is not False,
            "include_authors": raw_options.get("include_authors", True) is not False,
            "include_supported_versions": raw_options.get("include_supported_versions", True) is not False,
            "include_language_packs": bool(raw_options.get("include_language_packs", False)),
            "include_package_id": bool(raw_options.get("include_package_id", False)),
            "include_workshop_id": raw_options.get("include_workshop_id", True) is not False,
            "include_url": raw_options.get("include_url", True) is not False,
            "body_source": "description" if str(raw_options.get("body_source") or "").lower() == "description" else "notes",
            "image_name_source": "original" if str(raw_options.get("image_name_source") or "").lower() == "original" else "alias",
        }

    def _normalize_mod(self, index: int, mod: dict[str, Any], options: dict[str, Any]) -> dict[str, Any]:
        package_id = self._clean_inline_text(mod.get("package_id") or mod.get("package_id_raw"))
        raw_name = self._clean_inline_text(mod.get("name") or mod.get("display_name") or package_id or "未知模组")
        alias_name = self._clean_inline_text(mod.get("alias_name"))
        # 有别名时优先展示别名；没有别名就直接使用原名，避免导出里重复出现“名称/原名”。
        display_name = alias_name or self._clean_inline_text(mod.get("display_name")) or raw_name or package_id or "未知模组"
        body_value = mod.get("description") if options["body_source"] == "description" else mod.get("notes")
        group_names = self._normalize_string_list(mod.get("group_names") or [])
        # 作者在 Mod 数据里通常是列表；这里也兼容单个字符串，避免旧数据导出时丢失信息。
        authors = self._normalize_string_list(mod.get("author") or mod.get("authors") or [])
        tags = self._normalize_string_list(mod.get("tags") or [])
        supported_versions = self._normalize_string_list(mod.get("supported_versions") or [])
        return {
            "index": index,
            "sequence": f"{index:03d}",
            "package_id": package_id,
            "package_id_raw": self._clean_inline_text(mod.get("package_id_raw") or package_id),
            "display_name": display_name,
            "original_name": raw_name,
            "alias_name": alias_name,
            "has_alias": bool(alias_name and alias_name != raw_name),
            "body": self._clean_multiline_text(body_value),
            "tags": tags,
            "group_names": group_names,
            "authors": authors,
            "supported_versions": supported_versions,
            "language_packs": self._normalize_language_packs(mod.get("language_packs") or []),
            "workshop_id": self._clean_inline_text(mod.get("workshop_id")),
            "url": self._clean_inline_text(mod.get("url")),
            "preview_path": self._valid_image_path(mod.get("preview_path") or mod.get("cover_path") or mod.get("icon_path")),
        }

    def _normalize_string_list(self, values: Any) -> list[str]:
        source = values if isinstance(values, list) else [values]
        result = []
        seen = set()
        for value in source:
            text = self._clean_inline_text(value)
            if not text or text in seen:
                continue
            seen.add(text)
            result.append(text)
        return result

    def _normalize_language_packs(self, values: Any) -> list[dict[str, str]]:
        source = values if isinstance(values, list) else [values]
        result = []
        seen = set()
        for value in source:
            if not isinstance(value, dict):
                continue
            name = self._clean_inline_text(
                value.get("name") or value.get("original_name") or value.get("display_name") or value.get("package_id")
            )
            url = self._clean_inline_text(value.get("url"))
            if not name and not url:
                continue
            key = (name, url)
            if key in seen:
                continue
            seen.add(key)
            result.append({"name": name, "url": url})
        return result

    def _valid_image_path(self, path_value: Any) -> str:
        value = str(path_value or "").strip()
        if not value:
            return ""
        try:
            # 只接受当前磁盘上真实存在的图片路径，后续嵌入和复制就不需要反复处理不存在路径。
            path = Path(value)
            return str(path) if path.is_file() else ""
        except OSError:
            return ""

    def _clean_inline_text(self, value: Any) -> str:
        return re.sub(r"\s+", " ", self._clean_multiline_text(value)).strip()

    def _clean_multiline_text(self, value: Any) -> str:
        text = html.unescape(str(value or ""))
        # 备注和工坊描述可能带 HTML，导出介绍只保留用户能直接阅读的纯文本。
        text = re.sub(r"(?i)<br\s*/?>", "\n", text)
        text = re.sub(r"<[^>]+>", "", text)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = self._strip_existing_markdown_escapes(text)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
        return "\n".join(line for line in lines if line).strip()

    def _render_text_block(self, item: dict[str, Any], options: dict[str, Any]) -> str:
        if options["include_sequence"]:
            # 开启序号时标题已经包含名称，字段列表里不再重复输出“名称：”。
            lines = [f"{item['sequence']}. {item['display_name']}"]
            lines.extend(self._render_field_lines(item, options, include_name=False))
        else:
            lines = self._render_field_lines(item, options, include_name=True)
        return "\n".join(lines)

    def _build_export_fields(self, item: dict[str, Any], options: dict[str, Any], include_name: bool = True) -> list[dict[str, Any]]:
        # 所有格式共享同一份字段顺序，避免以后新增字段时各格式出现前后不一致。
        fields: list[dict[str, Any]] = []
        if include_name:
            fields.append({"label": "名称", "value": item["display_name"]})
        if item.get("has_alias") and item.get("original_name"):
            fields.append({"label": "原名", "value": item["original_name"]})
        if options["include_tags"] and item["tags"]:
            fields.append({"kind": "tags", "values": item["tags"]})
        if options["include_group_names"] and item["group_names"]:
            fields.append({"label": "分组", "value": "、".join(item["group_names"])})
        if options["include_authors"] and item["authors"]:
            fields.append({"label": "作者", "value": "、".join(item["authors"])})
        if options["include_supported_versions"] and item["supported_versions"]:
            fields.append({"label": "支持版本", "value": "、".join(item["supported_versions"])})
        fields.append({"label": "介绍", "value": item["body"] or "暂无介绍", "multiline": True})
        if options["include_package_id"] and item["package_id_raw"]:
            fields.append({"label": "包名", "value": item["package_id_raw"]})
        if options["include_workshop_id"] and item["workshop_id"]:
            fields.append({"label": "工坊ID", "value": item["workshop_id"]})
        if options["include_url"] and item["url"]:
            fields.append({"label": "网址", "value": item["url"], "url": True})
        if options["include_language_packs"]:
            for language_pack in item["language_packs"]:
                if language_pack["name"]:
                    fields.append({"label": "语言包", "value": language_pack["name"]})
                if language_pack["url"]:
                    fields.append({"label": "语言包网址", "value": language_pack["url"], "url": True})
        return fields

    def _render_field_lines(self, item: dict[str, Any], options: dict[str, Any], include_name: bool = True) -> list[str]:
        lines = []
        for field in self._build_export_fields(item, options, include_name):
            if field.get("kind") == "tags":
                lines.append(" ".join(f"#{tag}" for tag in field["values"]))
                continue
            lines.append(f"{field['label']}：{field['value']}")
        return lines

    def _render_markdown_field_lines(self, item: dict[str, Any], options: dict[str, Any], include_name: bool = True) -> list[str]:
        lines = []
        for field in self._build_export_fields(item, options, include_name):
            if field.get("kind") == "tags":
                # 标签保留 #tag 形式，标签内容本身仍按 Markdown 普通文本转义。
                lines.append(" ".join(f"#{self._escape_markdown_text(str(tag).lstrip('#'))}" for tag in field["values"]))
                continue
            value = field["value"]
            if field.get("url"):
                value = f"<{self._markdown_link_url(value)}>"
            elif field.get("multiline"):
                value = self._escape_markdown_multiline_text(value)
            else:
                value = self._escape_markdown_text(value)
            lines.append(f"{field['label']}：{value}")
        return lines

    def _write_text_file(self, target_path: str | None, source_name: str, items: list[dict[str, Any]], options: dict[str, Any]) -> dict[str, Any]:
        path = self._require_file_target(target_path)
        # TXT 和剪贴板共享纯文本渲染，保证两种导出的内容一致。
        path.write_text(self.render_plain_text(items, options, title=source_name), encoding="utf-8")
        return {"format": "txt", "path": str(path), "count": len(items)}

    def _write_markdown_dir(self, target_dir: str | None, source_name: str, items: list[dict[str, Any]], options: dict[str, Any]) -> dict[str, Any]:
        output_dir = self._require_dir_target(target_dir)
        image_dir = output_dir / "img"
        image_names: dict[int, str] = {}
        if options["include_cover"]:
            # Markdown 图片固定集中到 img 目录，便于把 md 文件和图片一起分享。
            image_dir.mkdir(parents=True, exist_ok=True)
            for item in items:
                image_name = self._markdown_image_name(item)
                if self._copy_cover_as_png(item["preview_path"], image_dir / image_name):
                    image_names[item["index"]] = image_name
        filename = self._unique_child_path(output_dir, f"{self.sanitize_filename(source_name, '模组推荐清单')}.md")
        filename.write_text(self.render_markdown(items, options, image_names, title=source_name), encoding="utf-8")
        return {
            "format": "markdown",
            "path": str(filename),
            "image_dir": str(image_dir) if image_names else "",
            "count": len(items),
            "image_count": len(image_names),
        }

    def _write_docx(self, target_path: str | None, source_name: str, items: list[dict[str, Any]], options: dict[str, Any]) -> dict[str, Any]:
        path = self._require_file_target(target_path)
        document = Document()
        self._configure_docx_styles(document)
        document.add_heading(source_name or "模组推荐清单", level=1)
        with tempfile.TemporaryDirectory(prefix="rmm_recommend_docx_") as temp_dir:
            for item in items:
                # DOCX 用标题承载名称，正文只放原名、标签、分组、作者和介绍等字段。
                title = item["display_name"]
                if options["include_sequence"]:
                    title = f"{item['sequence']}. {title}"
                document.add_heading(title, level=2)
                image_path = self._prepare_embed_image(item["preview_path"], temp_dir, item["sequence"])
                if options["include_cover"] and image_path:
                    document.add_picture(image_path, width=Inches(5.8))
                self._append_docx_fields(document, item, options)
                document.add_paragraph("")
        document.save(str(path))
        return {"format": "docx", "path": str(path), "count": len(items)}

    def _write_pdf(self, target_path: str | None, source_name: str, items: list[dict[str, Any]], options: dict[str, Any]) -> dict[str, Any]:
        path = self._require_file_target(target_path)
        font_name = self._register_pdf_font()
        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
            topMargin=16 * mm,
            bottomMargin=16 * mm,
        )
        story = []
        styles = self._build_pdf_styles(font_name)
        story.append(Paragraph(self._escape_pdf_text(source_name or "模组推荐清单"), styles["title"]))
        story.append(Spacer(1, 8))
        with tempfile.TemporaryDirectory(prefix="rmm_recommend_pdf_") as temp_dir:
            for item in items:
                # ReportLab 的 Paragraph 支持超链接标记，所以字段先格式化再写入 story。
                title = item["display_name"]
                if options["include_sequence"]:
                    title = f"{item['sequence']}. {title}"
                story.append(Paragraph(self._escape_pdf_text(title), styles["heading"]))
                image_path = self._prepare_embed_image(item["preview_path"], temp_dir, item["sequence"])
                if options["include_cover"] and image_path:
                    story.append(self._build_pdf_image(image_path, max_width=160 * mm, max_height=78 * mm))
                    story.append(Spacer(1, 6))
                for line in self._render_field_lines(item, options, include_name=False):
                    story.append(Paragraph(self._format_pdf_field_line(line).replace("\n", "<br/>"), styles["body"]))
                story.append(Spacer(1, 12))
            doc.build(story)
        return {"format": "pdf", "path": str(path), "count": len(items)}

    def _write_images(self, target_dir: str | None, source_name: str, items: list[dict[str, Any]], options: dict[str, Any]) -> dict[str, Any]:
        parent_dir = self._require_dir_target(target_dir)
        # 纯图片会先按导出名称创建独立文件夹，避免多次导出的单图混在同一个目录里。
        output_dir = self._unique_child_dir(parent_dir, self.sanitize_filename(source_name, "模组推荐图片"))
        output_dir.mkdir(parents=True, exist_ok=False)
        paths = []
        for item in items:
            # 纯图片是一模组一张图，重名时 _unique_child_path 会追加序号保护已有文件。
            filename = self._unique_child_path(output_dir, f"{self._image_filename_stem(item, options)}.png")
            self._render_recommendation_image(item, options, filename)
            paths.append(str(filename))
        return {"format": "image", "path": str(output_dir), "files": paths, "count": len(paths)}

    def _require_file_target(self, target_path: str | None) -> Path:
        path = Path(str(target_path or "").strip())
        if not str(path):
            raise ValueError("未指定导出文件路径")
        path.parent.mkdir(parents=True, exist_ok=True)
        return path

    def _require_dir_target(self, target_dir: str | None) -> Path:
        path = Path(str(target_dir or "").strip())
        if not str(path):
            raise ValueError("未指定导出目录")
        path.mkdir(parents=True, exist_ok=True)
        if not path.is_dir():
            raise ValueError("导出目录无效")
        return path

    def _unique_child_path(self, parent: Path, filename: str) -> Path:
        base_name = self.sanitize_filename(Path(filename).stem, "export")
        suffix = Path(filename).suffix
        candidate = parent / f"{base_name}{suffix}"
        index = 2
        # 不覆盖用户已有导出文件，重复导出时自动追加 _2、_3。
        while candidate.exists():
            candidate = parent / f"{base_name}_{index}{suffix}"
            index += 1
        return candidate

    def _unique_child_dir(self, parent: Path, dirname: str) -> Path:
        base_name = self.sanitize_filename(dirname, "export")
        candidate = parent / base_name
        index = 2
        # 目录也不覆盖已有内容，重复导出时每次创建新的结果文件夹。
        while candidate.exists():
            candidate = parent / f"{base_name}_{index}"
            index += 1
        return candidate

    def _markdown_image_name(self, item: dict[str, Any]) -> str:
        # Markdown 图片名按序号和包名固定，便于同一清单中的图片稳定引用。
        stem = self.sanitize_filename(f"{item['sequence']}-{item['package_id']}", item["sequence"])
        return f"{stem}.png"

    def _image_filename_stem(self, item: dict[str, Any], options: dict[str, Any]) -> str:
        # 纯图片文件名允许用户选择别名或原名，但前缀序号保持固定，方便按清单顺序浏览。
        name_source = item["original_name"] if options["image_name_source"] == "original" else item["display_name"]
        fallback = item["package_id"] or item["original_name"] or "模组"
        return self.sanitize_filename(f"{item['sequence']}-{name_source or fallback}", item["sequence"])

    def _copy_cover_as_png(self, source_path: str, target_path: Path) -> bool:
        if not source_path:
            return False
        try:
            with Image.open(source_path) as image:
                # 先转成 PNG 能统一 Markdown/DOCX/PDF 的图片输入格式，也顺便处理 EXIF 旋转。
                image = ImageOps.exif_transpose(image).convert("RGB")
                image.thumbnail((1600, 900), Image.Resampling.LANCZOS)
                image.save(target_path, "PNG", optimize=True)
            return True
        except Exception:
            try:
                shutil.copy2(source_path, target_path)
                return True
            except Exception:
                return False

    def _prepare_embed_image(self, source_path: str, temp_dir: str, sequence: str) -> str:
        if not source_path:
            return ""
        target_path = Path(temp_dir) / f"{sequence}.png"
        return str(target_path) if self._copy_cover_as_png(source_path, target_path) else ""

    def _configure_docx_styles(self, document: DocxDocument) -> None:
        font_name = self._document_font_name()
        for style_name in ("Normal", "Heading 1", "Heading 2"):
            style = document.styles[style_name]
            # 同时设置西文字体和东亚字体，避免 Word 里中文回退成不稳定的默认字体。
            style.font.name = font_name # type: ignore
            style.font.size = Pt(10.5 if style_name == "Normal" else 14) # type: ignore
            style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _append_docx_fields(self, document: DocxDocument, item: dict[str, Any], options: dict[str, Any]) -> None:
        for line in self._render_field_lines(item, options, include_name=False):
            label, separator, value = line.partition("：")
            paragraph = document.add_paragraph()
            if separator:
                # 字段名加粗，网址字段写成可点击链接，便于分享文档后直接打开来源。
                paragraph.add_run(f"{label}：").bold = True
                if label.endswith("网址") and value:
                    self._add_docx_hyperlink(paragraph, value)
                else:
                    paragraph.add_run(value)
            else:
                paragraph.add_run(line)

    def _add_docx_hyperlink(self, paragraph, url: str, text: str | None = None) -> None:
        # python-docx 没有公开的超链接 API，这里按 WordprocessingML 关系写入。
        relation_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relation_id)
        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(color)
        properties.append(underline)
        run.append(properties)
        text_node = OxmlElement("w:t")
        text_node.text = text or url
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _document_font_name(self) -> str:
        system = platform.system().lower()
        if system == "windows":
            return "Microsoft YaHei"
        if system == "darwin":
            return "PingFang SC"
        return "Noto Sans CJK SC"

    def _register_pdf_font(self) -> str:
        candidates = [
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf",
        ]
        for candidate in candidates:
            try:
                if Path(candidate).is_file():
                    # 优先使用系统中文字体，避免 PDF 导出后出现方块字。
                    pdfmetrics.registerFont(TTFont("RMMDocumentFont", candidate))
                    return "RMMDocumentFont"
            except Exception:
                continue
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        except Exception:
            pass
        return "STSong-Light"

    def _build_pdf_styles(self, font_name: str) -> dict[str, ParagraphStyle]:
        sample_styles = getSampleStyleSheet()
        return {
            "title": ParagraphStyle(
                "RMMRecommendationTitle",
                parent=sample_styles["Title"],
                fontName=font_name,
                fontSize=20,
                leading=26,
                textColor=colors.HexColor("#222222"),
                alignment=TA_LEFT,
            ),
            "heading": ParagraphStyle(
                "RMMRecommendationHeading",
                parent=sample_styles["Heading2"],
                fontName=font_name,
                fontSize=14,
                leading=20,
                spaceBefore=8,
                textColor=colors.HexColor("#222222"),
                wordWrap="CJK",
            ),
            "body": ParagraphStyle(
                "RMMRecommendationBody",
                parent=sample_styles["BodyText"],
                fontName=font_name,
                fontSize=10,
                leading=15,
                textColor=colors.HexColor("#333333"),
                wordWrap="CJK",
            ),
        }

    def _build_pdf_image(self, image_path: str, max_width: float, max_height: float) -> PdfImage:
        with Image.open(image_path) as image:
            width, height = image.size
        ratio = min(max_width / max(width, 1), max_height / max(height, 1))
        return PdfImage(image_path, width=width * ratio, height=height * ratio)

    def _render_recommendation_image(self, item: dict[str, Any], options: dict[str, Any], target_path: Path) -> None:
        width = 1080
        margin = 54
        cover_height = 560 if options["include_cover"] else 0
        title_font = self._load_font(36, bold=True)
        body_font = self._load_font(26)
        small_font = self._load_font(22)
        text_lines = self._build_image_text_lines(item, options, body_font, width - margin * 2)
        line_height = 40
        title_height = 58
        text_height = title_height + max(1, len(text_lines)) * line_height + margin * 2
        # 文本越多图片越高，但至少保留一个适合分享预览的基础高度。
        height = max(720, cover_height + text_height)
        animated_cover = self._build_animated_cover_frames(item["preview_path"], width, cover_height) if cover_height else None
        if animated_cover:
            cover_frames, durations, loop = animated_cover
            frames = [
                self._build_recommendation_image_frame(
                    cover, width, height, cover_height, margin, title_font, body_font, small_font, text_lines, line_height, item, options
                )
                for cover in cover_frames
            ]
            target_path.parent.mkdir(parents=True, exist_ok=True)
            frames[0].save(
                target_path,
                "PNG",
                save_all=True,
                append_images=frames[1:],
                duration=durations,
                loop=loop,
                optimize=True,
            )
            return

        cover = self._build_cover_image(item["preview_path"], width, cover_height) if cover_height else None
        canvas = self._build_recommendation_image_frame(
            cover, width, height, cover_height, margin, title_font, body_font, small_font, text_lines, line_height, item, options
        )
        target_path.parent.mkdir(parents=True, exist_ok=True)
        canvas.save(target_path, "PNG", optimize=True)

    def _build_recommendation_image_frame(
        self,
        cover: Image.Image | None,
        width: int,
        height: int,
        cover_height: int,
        margin: int,
        title_font: PillowFont,
        body_font: PillowFont,
        small_font: PillowFont,
        text_lines: list[str],
        line_height: int,
        item: dict[str, Any],
        options: dict[str, Any],
    ) -> Image.Image:
        canvas = Image.new("RGB", (width, height), "#f5f2ea")
        draw = ImageDraw.Draw(canvas)

        if cover_height and cover:
            canvas.paste(cover, (0, 0))

        text_top = cover_height
        draw.rectangle((0, text_top, width, height), fill="#f5f2ea")
        draw.rounded_rectangle((32, text_top + 28, width - 32, height - 34), radius=22, fill="#ffffff", outline="#ddd6c8", width=2)
        title = item["display_name"]
        if options["include_sequence"]:
            title = f"{item['sequence']}. {title}"
        draw.text((margin, text_top + 52), title, font=title_font, fill="#202020")

        y = text_top + 112
        for line in text_lines:
            fill = "#666666" if line.startswith("#") else "#2d2d2d"
            draw.text((margin, y), line, font=small_font if line.startswith("#") else body_font, fill=fill)
            y += line_height
        return canvas

    def _build_cover_image(self, source_path: str, width: int, height: int) -> Image.Image:
        if not source_path:
            # 缺封面时生成简单占位图，保证纯图片导出仍有稳定版式。
            cover = Image.new("RGB", (width, height), "#d8d1c4")
            draw = ImageDraw.Draw(cover)
            font = self._load_font(32, bold=True)
            text = "暂无封面"
            bbox = draw.textbbox((0, 0), text, font=font)
            draw.text(((width - (bbox[2] - bbox[0])) / 2, (height - (bbox[3] - bbox[1])) / 2), text, font=font, fill="#726b60")
            return cover
        try:
            with Image.open(source_path) as image:
                image = ImageOps.exif_transpose(image).convert("RGB")
                # 封面按固定区域居中裁切，避免不同尺寸封面撑乱介绍图版式。
                return ImageOps.fit(image, (width, height), method=Image.Resampling.LANCZOS, centering=(0.5, 0.5))
        except Exception:
            return self._build_cover_image("", width, height)

    def _build_animated_cover_frames(self, source_path: str, width: int, height: int) -> tuple[list[Image.Image], list[int], int] | None:
        if not source_path:
            return None
        try:
            with Image.open(source_path) as image:
                if getattr(image, "n_frames", 1) <= 1:
                    return None
                frames = []
                durations = []
                loop = int(image.info.get("loop", 0) or 0)
                for frame_index in range(image.n_frames):
                    image.seek(frame_index)
                    frame = ImageOps.exif_transpose(image.copy()).convert("RGB")
                    frames.append(ImageOps.fit(frame, (width, height), method=Image.Resampling.LANCZOS, centering=(0.5, 0.5)))
                    duration = frame.info.get("duration", image.info.get("duration", 100))
                    durations.append(max(20, int(duration or 100)))
                return (frames, durations, loop) if frames else None
        except Exception:
            return None

    def _build_image_text_lines(self, item: dict[str, Any], options: dict[str, Any], font: PillowFont, max_width: int) -> list[str]:
        lines = []
        for field_line in self._render_field_lines(item, options, include_name=False):
            label, separator, value = field_line.partition("：")
            raw_lines = [field_line] if not separator else [f"{label}：{value}"]
            for raw_line in raw_lines:
                lines.extend(self._wrap_text(raw_line, font, max_width))
        return lines

    def _wrap_text(self, text: str, font: PillowFont, max_width: int) -> list[str]:
        result = []
        draw = ImageDraw.Draw(Image.new("RGB", (1, 1)))
        for source_line in str(text or "").split("\n"):
            current = ""
            for char in source_line:
                candidate = f"{current}{char}"
                bbox = draw.textbbox((0, 0), candidate, font=font)
                if current and bbox[2] - bbox[0] > max_width:
                    # 中英文混排很难按单词稳定折行，按字符测量能保证文字不越出图片边界。
                    result.append(current)
                    current = char
                else:
                    current = candidate
            if current:
                result.append(current)
        return result or [""]

    def _load_font(self, size: int, bold: bool = False) -> PillowFont:
        candidates = [
            "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "/System/Library/Fonts/PingFang.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]
        for candidate in candidates:
            try:
                if Path(candidate).is_file():
                    return ImageFont.truetype(candidate, size)
            except Exception:
                continue
        return ImageFont.load_default()

    def _strip_existing_markdown_escapes(self, value: Any) -> str:
        return self.EXISTING_MARKDOWN_ESCAPE_RE.sub(r"\1", str(value or ""))

    def _markdown_alt_text(self, value: Any) -> str:
        text = self._clean_inline_text(value)
        # alt 文本不需要 Markdown 方括号语法，替换成括号可以避免图片语法被截断。
        text = text.replace("[", "(").replace("]", ")")
        text = re.sub(r"\s+", " ", text).strip()
        return text or "封面"

    def _markdown_link_url(self, value: Any) -> str:
        text = self._strip_existing_markdown_escapes(value).strip()
        return text.replace(" ", "%20").replace("<", "%3C").replace(">", "%3E")

    def _escape_markdown_text(self, value: Any) -> str:
        text = self._strip_existing_markdown_escapes(value)
        return self.MARKDOWN_ESCAPE_RE.sub(r"\\\1", text)

    def _escape_markdown_multiline_text(self, value: Any) -> str:
        text = self._strip_existing_markdown_escapes(value)
        return "\n".join(self._escape_markdown_text(line) for line in str(text or "").split("\n"))

    def _format_pdf_field_line(self, line: str) -> str:
        label, separator, value = str(line or "").partition("：")
        if separator and label.endswith("网址") and value:
            escaped_label = self._escape_pdf_text(label)
            escaped_href = html.escape(value, quote=True)
            escaped_text = self._escape_pdf_text(value)
            return f'{escaped_label}：<a href="{escaped_href}">{escaped_text}</a>'
        return self._escape_pdf_text(line)

    def _escape_pdf_text(self, value: Any) -> str:
        return html.escape(str(value or ""))
